import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_header_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = "Tensor Decomposer — SE801 Midterm Technical Report  |  Page "
        # Add page number XML element
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        p._p.append(fldSimple)

def build_report():
    doc = Document()

    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Dark Charcoal
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Colors
    NAVY = RGBColor(0x0F, 0x17, 0x2A)
    TEAL = RGBColor(0x02, 0x84, 0xC7)
    SLATE = RGBColor(0x47, 0x55, 0x69)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = TEAL
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = SLATE
        return p

    def add_callout_code(text_code, title="Editable Draw.io / Mermaid Source Code"):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        p_title.paragraph_format.space_after = Pt(2)
        run_t = p_title.add_run(f"[{title}]")
        run_t.font.size = Pt(9.5)
        run_t.font.bold = True
        run_t.font.color.rgb = TEAL

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        # Border setup
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>'
            f'  <w:top w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text_code)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------------------------
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(36)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_main = p_title.add_run("TENSOR DECOMPOSER\n")
    run_main.font.name = 'Arial'
    run_main.font.size = Pt(26)
    run_main.font.bold = True
    run_main.font.color.rgb = NAVY

    run_sub = p_title.add_run("An Interactive High-Dimensional Multiway Data Analysis and Benchmarking System")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = TEAL

    p_rep = doc.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_before = Pt(30)
    run_rep = p_rep.add_run("MIDTERM TECHNICAL REPORT\nSE 801 — Software Project Lab III")
    run_rep.font.name = 'Arial'
    run_rep.font.size = Pt(13)
    run_rep.font.bold = True
    run_rep.font.color.rgb = SLATE

    # Divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_before = Pt(20)
    p_div.paragraph_format.space_after = Pt(40)
    run_div = p_div.add_run("―" * 35)
    run_div.font.color.rgb = TEAL
    run_div.font.bold = True

    # Details table
    tbl_details = doc.add_table(rows=3, cols=2)
    tbl_details.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    details_data = [
        ("Submitted By:", "Md. Kamrul Hasan\nRoll: 1418 | Phone: 01535757066\nBSSE 8th Semester\nInstitute of Information Technology\nUniversity of Dhaka"),
        ("Supervised By:", "Mohd. Zulfiquar Hafiz\nAssistant Professor\nInstitute of Information Technology\nUniversity of Dhaka"),
        ("Submitted To:", "SPL 3 Program Committee\nInstitute of Information Technology\nUniversity of Dhaka\nDate of Submission: August 2026")
    ]

    for idx, (label, val) in enumerate(details_data):
        row = tbl_details.rows[idx]
        cell_l, cell_r = row.cells[0], row.cells[1]
        cell_l.width = Inches(2.2)
        cell_r.width = Inches(4.3)
        
        p_l = cell_l.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.font.bold = True
        r_l.font.size = Pt(10.5)
        r_l.font.color.rgb = NAVY
        
        p_r = cell_r.paragraphs[0]
        r_r = p_r.add_run(val)
        r_r.font.size = Pt(10)
        r_r.font.color.rgb = SLATE

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # LETTER OF TRANSMITTAL
    # ---------------------------------------------------------------------------
    add_heading_1("Letter of Transmittal")
    p_date = doc.add_paragraph("Date: August 2026")
    p_date.runs[0].font.bold = True
    
    p_to = doc.add_paragraph(
        "To,\n"
        "The SPL 3 Program Committee & Exam Committee\n"
        "Institute of Information Technology (IIT)\n"
        "University of Dhaka, Dhaka-1000\n"
    )
    p_to.runs[0].font.bold = True

    p_sub = doc.add_paragraph("Subject: Submission of Midterm Technical Report on 'Tensor Decomposer'")
    p_sub.runs[0].font.bold = True

    doc.add_paragraph(
        "Dear Sir/Madam,\n\n"
        "I am pleased to submit the Midterm Technical Report for my Software Project Lab III (SE 801) project, titled "
        "\"Tensor Decomposer: An Interactive High-Dimensional Multiway Data Analysis and Benchmarking System\".\n\n"
        "This report details the requirements analysis, system architecture, scenario-based system modeling, activity workflows, "
        "preliminary test plans, and implementation roadmap for the web-based tensor decomposition platform. The system incorporates "
        "higher-order multilinear algebraic algorithms—including CP Decomposition, Tucker Decomposition, HOSVD, and Tensor Train Decomposition, "
        "along with standard matrix factorizations—paired with single-run FLOPs benchmarking, error metrics analysis, and dynamic 3D visual representations.\n\n"
        "I express my sincere gratitude for your guidance and continuous support throughout the project. I look forward to your valuable feedback and recommendations."
    )

    p_sign = doc.add_paragraph("\nSincerely,\n\n_______________________\nMd. Kamrul Hasan\nRoll: 1418\nBSSE 8th Semester, IIT, University of Dhaka")
    p_sign.runs[0].font.bold = True

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # ACKNOWLEDGEMENT & ABSTRACT
    # ---------------------------------------------------------------------------
    add_heading_1("Acknowledgement")
    doc.add_paragraph(
        "First and foremost, I express my deepest gratitude to Almighty for providing me with the strength, endurance, and capability "
        "to complete the midterm milestones of this project successfully.\n\n"
        "I extend my sincere appreciation to my project supervisor, Mohd. Zulfiquar Hafiz, Assistant Professor, Institute of Information Technology (IIT), "
        "University of Dhaka, for his invaluable guidance, scholarly advice, and constructive reviews throughout the formulation and implementation "
        "of this research project.\n\n"
        "I am also profoundly grateful to the faculty members and SPL 3 Program Committee of the Institute of Information Technology for providing a rich "
        "academic environment and technical infrastructure that made this work possible."
    )

    add_heading_1("Abstract")
    doc.add_paragraph(
        "Modern datasets across scientific disciplines—such as deep learning, image processing, signal processing, neuroimaging, and chemometrics—are "
        "inherently multi-dimensional and naturally structured as higher-order tensors. Traditional matrix factorization methods like SVD or PCA require "
        "flattening (matricizing) multiway data, which severely destroys higher-order spatial and multiway structural correlations. Conversely, computing "
        "higher-order tensor factorizations introduces substantial computational complexity, memory overhead, and parameter explosion.\n\n"
        "To bridge this gap, this report presents Tensor Decomposer, an interactive, web-based analytical platform designed to ingest, decompose, benchmark, "
        "and visualize higher-dimensional tensor datasets. The platform implements core tensor decomposition algorithms—specifically Canonical Polyadic (CP) "
        "Decomposition via Alternating Least Squares (CP-ALS), Tucker Decomposition via Higher-Order Orthogonal Iteration (HOOI/Tucker-ALS), Higher-Order "
        "Singular Value Decomposition (HOSVD), and Tensor Train (TT) Decomposition—alongside standard matrix factorizations (SVD, QR, LU, Eigendecomposition). "
        "The system incorporates high-precision single-run performance analytics (FLOPs estimation, computational complexity formulas, compression ratios) "
        "and reconstruction accuracy evaluations (MAE, RMSE, Relative Error). Furthermore, an interactive SVG/isometric visual engine renders 3D decomposition "
        "structures, component spectra, and matrix heatmaps. This report outlines the system requirements, scenario-based modeling, activity workflows, "
        "preliminary test plans, and implementation progress."
    )

    add_heading_1("Table of Contents")
    toc_items = [
        ("Letter of Transmittal", "2"),
        ("Acknowledgement", "3"),
        ("Abstract", "3"),
        ("1. Project Overview", "5"),
        ("   1.1 Project Title", "5"),
        ("   1.2 Problem Statement", "5"),
        ("   1.3 Objectives", "6"),
        ("   1.4 Scope", "6"),
        ("   1.5 Deliverables", "7"),
        ("2. Requirements Analysis", "8"),
        ("   2.1 Functional Requirements", "8"),
        ("   2.2 Non-Functional Requirements", "10"),
        ("   2.3 Stakeholders", "11"),
        ("3. System Modeling", "12"),
        ("   3.1 Use Case Diagram & Descriptions", "12"),
        ("   3.2 Detailed Use Case Specifications", "14"),
        ("   3.3 Activity Diagram & Description", "18"),
        ("7. Preliminary Test Plan", "21"),
        ("   7.1 Testing Objectives", "21"),
        ("   7.2 Features to be Tested", "21"),
        ("8. Timeline", "25"),
        ("   8.1 Project Gantt Chart", "25"),
        ("   8.2 Work Breakdown Schedule", "26"),
        ("9. Conclusion & Future Work", "27"),
        ("10. References", "28")
    ]

    tbl_toc = doc.add_table(rows=len(toc_items), cols=2)
    tbl_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (item, pg) in enumerate(toc_items):
        r = tbl_toc.rows[idx]
        c0, c1 = r.cells[0], r.cells[1]
        c0.width = Inches(5.5)
        c1.width = Inches(1.0)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(item)
        if not item.startswith("   "):
            r0.font.bold = True
            r0.font.color.rgb = NAVY
        else:
            r0.font.color.rgb = SLATE
            
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(pg)
        r1.font.bold = True
        r1.font.color.rgb = TEAL

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # SECTION 1: PROJECT OVERVIEW
    # ---------------------------------------------------------------------------
    add_heading_1("1. Project Overview")
    
    add_heading_2("1.1 Project Title")
    doc.add_paragraph("Tensor Decomposer: An Interactive High-Dimensional Multiway Data Analysis and Benchmarking System")

    add_heading_2("1.2 Problem Statement")
    doc.add_paragraph(
        "In modern data-driven fields such as computer vision, deep learning, signal processing, neuroimaging (EEG/fMRI), and chemometrics, "
        "data naturally manifests in multi-dimensional formats (tensors) rather than traditional 2D matrices. For example, color video sequences "
        "possess spatial, temporal, and chromatic dimensions (Height x Width x Frames x Channels).\n\n"
        "When conventional matrix reduction techniques like Principal Component Analysis (PCA) or Singular Value Decomposition (SVD) are applied to "
        "higher-order data, the tensor must first be flattened into a 2D matrix. This matricization severely destroys intrinsic multiway correlations, "
        "spatial patterns, and multi-linear structure. Conversely, operating directly on raw high-dimensional tensors suffers from the 'curse of dimensionality', "
        "leading to exponential growth in computational memory, execution time, and storage overhead.\n\n"
        "While tensor decomposition algorithms—such as Canonical Polyadic (CP) decomposition, Tucker decomposition, Higher-Order SVD (HOSVD), and Tensor Train (TT) "
        "decomposition—offer powerful mathematical frameworks for multiway data compression and feature extraction, existing software solutions suffer from several critical shortcomings:\n"
        "1. Fragmented Implementations: Algorithms are scattered across disparate language ecosystems (MATLAB Tensor Toolbox, Python TensorLy, C++ packages) without a unified interface.\n"
        "2. Absence of Real-Time Analytics: Existing tools lack automated single-run computational benchmarking (FLOPs estimation, complexity formulas, compression ratios).\n"
        "3. Lack of Interactive Visualizations: Researchers lack web-accessible visual representations to inspect decomposed core tensors, component factor spectra, and 3D equation layouts dynamically."
    )

    add_heading_2("1.3 Objectives")
    doc.add_paragraph(
        "The primary goal of the Tensor Decomposer project is to architect and deliver a unified web-based platform for interactive tensor decomposition, "
        "reconstruction accuracy analysis, performance benchmarking, and visualization. The specific technical objectives include:\n"
        "• Comprehensive Algorithm Suite: Implement robust mathematical algorithms for tensor decompositions (CP-ALS, HOOI/Tucker-ALS, HOSVD, Tensor Train) and core matrix factorizations (SVD, QR, LU, Eigendecomposition).\n"
        "• In-Browser Interactive Interface: Develop a responsive frontend capable of ingesting raw manual numerical tensors or uploaded data files (JSON/CSV) and rendering immediate analytical output.\n"
        "• Precision Reconstruction & Error Metrics: Provide exact numerical evaluation of reconstruction quality by calculating Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), Relative Error, and reconstructed element head previews.\n"
        "• Performance Benchmarking Engine: Compute single-run execution times (ms), floating-point operations (FLOPs count), theoretical complexity formulas, parameter compression counts, and compression ratios.\n"
        "• Multi-Method Comparative Benchmarking: Enable side-by-side algorithmic comparisons across compression efficiency, reconstruction error, and computational runtime.\n"
        "• 3D Isometric & Spectral Visualizations: Render scaled isometric 3D block diagrams, eigenvalue/singular value spectra, and matrix heatmaps dynamically using SVG."
    )

    add_heading_2("1.4 Scope")
    doc.add_paragraph(
        "The project encompasses the design, algorithm development, frontend visual engine, and backend web integration for the Tensor Decomposer application.\n\n"
        "In Scope:\n"
        "• Web application built with Django (backend service orchestration) and modern ES Modules JavaScript (frontend interactivity).\n"
        "• Implementation of CP-ALS, Tucker-ALS (HOOI), HOSVD, Tensor Train, SVD, QR, LU, and Eigendecomposition algorithms.\n"
        "• Interactive UI components featuring theme toggling (Dark/Light mode), mutual-exclusion inputs (manual text vs. file upload), and AJAX loading states.\n"
        "• Calculation of compression ratios, parameter counts, FLOPs estimates, MAE, RMSE, relative errors, and preview heads.\n"
        "• Dynamic generation of isometric SVG equations, bar chart spectra, and scalable matrix heatmaps.\n\n"
        "Out of Scope (for current midterm release):\n"
        "• Distributed multi-GPU CUDA cluster execution for terabyte-scale tensors.\n"
        "• Real-time streaming tensor decomposition for continuous live sensor data."
    )

    add_heading_2("1.5 Deliverables")
    doc.add_paragraph(
        "1. Complete Web Application Codebase: Django backend and modular ES Module frontend static JavaScript library (`home.js`, `theme.js`, `helpers.js`, `charts.js`, `visualizer.js`).\n"
        "2. Core Mathematical Algorithm Suite: Fully tested Python modules for CP, Tucker, HOSVD, TT, SVD, QR, LU, and Eigendecomposition.\n"
        "3. Interactive Visualization Engine: Client-side SVG graphics generator for 3D isometric equation layouts, spectra, and heatmaps.\n"
        "4. Midterm Technical Report & System Documentation: Comprehensive software architecture, requirements analysis, system modeling, test plans, and project timeline."
    )

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # SECTION 2: REQUIREMENTS ANALYSIS
    # ---------------------------------------------------------------------------
    add_heading_1("2. Requirements Analysis")

    add_heading_2("2.1 Functional Requirements")
    doc.add_paragraph("The functional requirements of the Tensor Decomposer platform specify the core capabilities and services provided by the application.")

    fr_data = [
        ("FR-01", "Manual Tensor Input", "System shall allow users to manually enter tensor data in array format (1D, 2D, 3D, 4D)."),
        ("FR-02", "File Upload Ingestion", "System shall accept file uploads containing tensor structures formatted in JSON or CSV."),
        ("FR-03", "Mutual Exclusion Input", "System shall automatically clear file selection upon manual text edit, and vice-versa."),
        ("FR-04", "Algorithm Selection", "System shall allow selection among CP, Tucker, HOSVD, TT, SVD, QR, LU, and Eigendecomposition."),
        ("FR-05", "CP Decomposition (CP-ALS)", "System shall compute rank-R CANDECOMP/PARAFAC factor matrices and weight vectors using CP-ALS."),
        ("FR-06", "Tucker Decomposition (HOOI)", "System shall compute core tensor and orthogonal factor matrices using Higher-Order Orthogonal Iteration."),
        ("FR-07", "Tensor Train Decomposition", "System shall compute sequence of 3D core tensors connecting tensor modes."),
        ("FR-08", "Matrix Factorization Suite", "System shall perform SVD, QR, LU, and Eigendecomposition for 2D matrix inputs."),
        ("FR-09", "Reconstruction Analysis", "System shall reconstruct original tensor from decomposed factors and measure absolute & relative errors."),
        ("FR-10", "Accuracy Metrics", "System shall calculate Mean Absolute Error (MAE) and Root Mean Squared Error (RMSE)."),
        ("FR-11", "FLOPs Benchmarking", "System shall estimate total Floating Point Operations (FLOPs) and format in KFLOPs/MFLOPs/GFLOPs."),
        ("FR-12", "Single-Run Performance", "System shall measure execution time (ms), compressed parameter counts, and compression ratios."),
        ("FR-13", "Algorithmic Comparison", "System shall run multi-method benchmarking comparisons across all algorithms side-by-side."),
        ("FR-14", "3D Isometric Visualization", "System shall generate 3D isometric equation blocks scaling based on tensor shapes."),
        ("FR-15", "Theme Management", "System shall support theme toggling (Dark/Light mode) persisting across browser sessions.")
    ]

    tbl_fr = doc.add_table(rows=len(fr_data)+1, cols=3)
    tbl_fr.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_fr)

    headers_fr = ["Req ID", "Feature Name", "Description & Functional Specification"]
    hdr_cells = tbl_fr.rows[0].cells
    for i, title in enumerate(headers_fr):
        hdr_cells[i].width = Inches(1.0 if i==0 else (1.8 if i==1 else 3.7))
        set_cell_background(hdr_cells[i], "0F172A")
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (rid, fname, fdesc) in enumerate(fr_data):
        row_cells = tbl_fr.rows[idx+1].cells
        if idx % 2 == 1:
            for c in row_cells: set_cell_background(c, "F8FAFC")
        
        p0 = row_cells[0].paragraphs[0]
        r0 = p0.add_run(rid)
        r0.font.bold = True
        r0.font.color.rgb = TEAL

        p1 = row_cells[1].paragraphs[0]
        r1 = p1.add_run(fname)
        r1.font.bold = True
        
        p2 = row_cells[2].paragraphs[0]
        p2.add_run(fdesc)

    add_heading_2("2.2 Non-Functional Requirements")
    nfr_data = [
        ("NFR-01", "Performance", "Algorithm execution for small to medium tensors (< 100x100x100) shall return results within < 500 ms."),
        ("NFR-02", "Reliability & Accuracy", "Matrix and tensor decomposition algorithms shall maintain floating point accuracy within 1e-6 error bounds."),
        ("NFR-03", "Usability & Responsive UI", "The user interface shall be intuitive, mobile-responsive, and maintain high-contrast dark/light themes."),
        ("NFR-04", "Maintainability & Quality", "Frontend code shall follow modular ES Module standards, separating DOM logic, formatting, and SVG rendering."),
        ("NFR-05", "Portability & Compatibility", "System shall function across modern web browsers (Chrome, Brave, Firefox, Edge, Safari) without external plugins."),
        ("NFR-06", "Error Handling & Feedback", "Invalid user inputs (e.g., non-numeric arrays, mismatched ranks) shall display informative error panels."),
        ("NFR-07", "Security & Privacy", "Uploaded tensor dataset files shall be processed in-memory or in isolated temporary storage without public access."),
        ("NFR-08", "Scalability", "Backend architecture shall support multi-threaded asynchronous processing of concurrent user requests.")
    ]

    tbl_nfr = doc.add_table(rows=len(nfr_data)+1, cols=3)
    tbl_nfr.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_nfr)

    hdr_nfr = tbl_nfr.rows[0].cells
    for i, title in enumerate(["NFR ID", "Category", "Quality Requirement Specification"]):
        hdr_nfr[i].width = Inches(1.0 if i==0 else (1.8 if i==1 else 3.7))
        set_cell_background(hdr_nfr[i], "0284C7")
        p = hdr_nfr[i].paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (nid, ncat, ndesc) in enumerate(nfr_data):
        row_cells = tbl_nfr.rows[idx+1].cells
        if idx % 2 == 1:
            for c in row_cells: set_cell_background(c, "F8FAFC")
        
        row_cells[0].paragraphs[0].add_run(nid).font.bold = True
        row_cells[1].paragraphs[0].add_run(ncat).font.bold = True
        row_cells[2].paragraphs[0].add_run(ndesc)

    add_heading_2("2.3 Stakeholders")
    doc.add_paragraph(
        "1. Academic Researchers & Data Scientists: Primary users who utilize the system to compress higher-dimensional tensors, benchmark decomposition runtimes, and inspect component factor matrices.\n"
        "2. Students & Educators: Users seeking interactive visual tools to understand multi-linear algebra concepts (CP vs. Tucker vs. HOSVD).\n"
        "3. System Developers & Maintainers: Engineers responsible for expanding the algorithm suite, maintaining backend Python services, and optimizing web frontend components.\n"
        "4. Project Supervisors & Academic Reviewers: Evaluators monitoring system adherence to software engineering standards, theoretical correctness, and performance benchmarks."
    )

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # SECTION 3: SYSTEM MODELING
    # ---------------------------------------------------------------------------
    add_heading_1("3. System Modeling")

    add_heading_2("3.1 Use Case Diagram & Descriptions")
    doc.add_paragraph(
        "The Use Case Diagram illustrates the functional interactions between the User (Researcher/Data Scientist) and the Tensor Decomposer system boundary."
    )

    # Insert Use Case Image
    if os.path.exists("report_assets/use_case_diagram.png"):
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img = doc.paragraphs[-1]
        p_img.add_run().add_picture("report_assets/use_case_diagram.png", width=Inches(5.8))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figure 3.1: Tensor Decomposer Use Case Diagram")
        r_cap.font.bold = True
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = SLATE

    # Draw.io / Mermaid code callout
    use_case_mermaid = """graph LR
    User((User / Researcher)) --> UC1[UC-1: Input / Upload Tensor Data]
    User --> UC2[UC-2: Select Decomposition Algorithm]
    User --> UC3[UC-3: Execute Decomposition]
    User --> UC4[UC-4: Analyze Accuracy & Errors]
    User --> UC5[UC-5: Benchmark FLOPs & Speed]
    User --> UC6[UC-6: Compare Methods]
    User --> UC7[UC-7: Render 3D Visuals & Heatmaps]

    subgraph Tensor Decomposer System Boundary
        UC1
        UC2
        UC3
        UC4
        UC5
        UC6
        UC7
    end"""
    add_callout_code(use_case_mermaid, "Editable Draw.io / Mermaid Source Code — Use Case Diagram")

    add_heading_2("3.2 Detailed Use Case Specifications")
    doc.add_paragraph("Below are the detailed formal specifications for the primary use cases of the Tensor Decomposer system.")

    uc_specs = [
        ("UC-1: Ingest & Validate Tensor Input", "User / Researcher", "Web Interface displays input form.", "Tensor array is validated and stored in system state.", 
         "1. User accesses application home page.\n2. User enters array text manually OR selects a JSON/CSV file.\n3. System validates array dimensions and numeric formatting.\n4. System clears opposite input field (mutual exclusion)."),
        
        ("UC-2: Execute Tensor Decomposition", "User / Researcher", "Valid tensor input is loaded into state.", "Decomposed factor matrices/cores are generated and rendered.",
         "1. User selects target algorithm (CP, Tucker, HOSVD, TT, SVD, QR, LU).\n2. User clicks 'Decompose' button.\n3. Backend executes algorithm (CP-ALS / HOOI).\n4. System returns formatted factor matrices and core tensors."),

        ("UC-3: Analyze Reconstruction Accuracy", "User / Researcher", "Decomposition execution has completed.", "MAE, RMSE, Relative Error, and reconstructed head preview are displayed.",
         "1. User clicks 'Analyze' action.\n2. Backend reconstructs tensor from factor matrices.\n3. System computes MAE, RMSE, and Relative Error.\n4. System displays accuracy summary panel with top reconstructed elements."),

        ("UC-4: Benchmark Single-Run Performance", "User / Researcher", "Algorithm and tensor input are configured.", "Performance analytics (FLOPs, runtime, compression ratio) rendered.",
         "1. User triggers 'Benchmark' action.\n2. System measures execution runtime over iterations.\n3. Backend calculates FLOPs count and parameter reduction ratio.\n4. System populates benchmark analytics panel with metrics."),

        ("UC-5: Compare Decomposition Methods", "User / Researcher", "Tensor input is uploaded/entered.", "Comparative summary table and multi-bar benchmark charts displayed.",
         "1. User triggers 'Compare' action.\n2. System sequentially executes all supported algorithms.\n3. Backend compiles side-by-side table of runtimes, errors, and compression ratios.\n4. Frontend renders comparative visualization charts."),

        ("UC-6: Render Interactive Visualizations", "User / Researcher", "Decomposition result data is available.", "3D isometric equation diagram, spectra, and heatmaps rendered.",
         "1. User clicks 'Visualize' action.\n2. Frontend reads tensor shapes and factor matrices.\n3. SVG Engine draws 3D scaled isometric blocks for equation structure.\n4. Heatmap generator draws normalized 2D/3D slice matrix heatmaps.")
    ]

    for u_title, u_actor, u_pre, u_post, u_flow in uc_specs:
        add_heading_3(u_title)
        tbl_uc = doc.add_table(rows=5, cols=2)
        tbl_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_uc)
        
        u_data = [
            ("Primary Actor", u_actor),
            ("Pre-Conditions", u_pre),
            ("Post-Conditions", u_post),
            ("Main Success Scenario", u_flow)
        ]
        
        c_hdr = tbl_uc.rows[0].cells
        c_hdr[0].width = Inches(1.8)
        c_hdr[1].width = Inches(4.7)
        set_cell_background(c_hdr[0], "0F172A")
        set_cell_background(c_hdr[1], "0F172A")
        r_h0 = c_hdr[0].paragraphs[0].add_run("Use Case Attribute")
        r_h0.font.bold = True
        r_h0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h1 = c_hdr[1].paragraphs[0].add_run("Specification Detail")
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i_row, (k, v) in enumerate(u_data):
            r_c = tbl_uc.rows[i_row+1].cells
            r_c[0].width = Inches(1.8)
            r_c[1].width = Inches(4.7)
            if i_row % 2 == 1:
                set_cell_background(r_c[0], "F8FAFC")
                set_cell_background(r_c[1], "F8FAFC")
            
            p_k = r_c[0].paragraphs[0].add_run(k)
            p_k.font.bold = True
            p_k.font.color.rgb = TEAL
            r_c[1].paragraphs[0].add_run(v)

    add_heading_2("3.3 Activity Diagram & Description")
    doc.add_paragraph(
        "The Activity Diagram models the dynamic operational workflow of the system from user tensor submission through validation, "
        "algorithmic execution, error analysis, benchmarking, and interactive SVG visualization rendering."
    )

    # Insert Activity Diagram Image
    if os.path.exists("report_assets/activity_diagram.png"):
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img = doc.paragraphs[-1]
        p_img.add_run().add_picture("report_assets/activity_diagram.png", width=Inches(4.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figure 3.2: Tensor Decomposer System Activity Workflow")
        r_cap.font.bold = True
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = SLATE

    # Draw.io / Mermaid code callout
    act_mermaid = """flowchart TD
    Start([Start]) --> Input[User Submits Tensor Input / File Upload]
    Input --> Validate{Is Tensor Input Valid?}
    Validate -- No --> Error[Display Error Message Panel]
    Error --> Input
    Validate -- Yes --> Exec[Execute Selected Decomposition Algorithm]
    Exec --> Recon[Reconstruct Tensor & Calculate Error Metrics]
    Recon --> Bench[Calculate FLOPs, Complexity & Compression Ratios]
    Bench --> Format[Format Matrix Outputs to 4 Decimals]
    Format --> Render[Render 3D Isometric Equations & Heatmaps]
    Render --> Stop([End])"""
    add_callout_code(act_mermaid, "Editable Draw.io / Mermaid Source Code — Activity Diagram")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # SECTION 7: PRELIMINARY TEST PLAN
    # ---------------------------------------------------------------------------
    add_heading_1("7. Preliminary Test Plan")

    add_heading_2("7.1 Testing Objectives")
    doc.add_paragraph(
        "The primary objectives of the preliminary testing phase are:\n"
        "1. Algorithmic Precision Verification: Ensure matrix and tensor factorizations (CP-ALS, HOOI, HOSVD, TT, SVD, QR, LU) yield numerical solutions matching mathematical ground truth within 1e-6 error margins.\n"
        "2. Input Validation Robustness: Verify that invalid array formats, malformed JSON files, mismatched matrix dimensions, and empty inputs trigger clean user-facing error panels without application crash.\n"
        "3. Performance Analytics Validation: Confirm that single-run FLOPs estimations, compression ratio calculations, and execution duration measurements are consistent and accurate.\n"
        "4. Frontend Interactivity & UI Stability: Ensure theme toggling, mutual exclusion input handling, loading skeleton overlays, and dynamic SVG visual cards render reliably across target browsers."
    )

    add_heading_2("7.2 Features to be Tested")
    doc.add_paragraph("The preliminary high-level test suite specifies test scenarios, inputs, expected outputs, and acceptance criteria across all system modules.")

    test_cases = [
        ("TC-01", "Manual Tensor Validation", "Enter valid 3D array [[[1,2],[3,4]],[[5,6],[7,8]]].", "System accepts input, parses 2x2x2 shape.", "PASS"),
        ("TC-02", "Invalid Text Handling", "Enter malformed string '[[1, 2], [invalid]]'.", "System catches syntax error, displays clear error panel.", "PASS"),
        ("TC-03", "Mutual Exclusion Input", "Select file upload while text field contains tensor data.", "Text field automatically clears upon file upload.", "PASS"),
        ("TC-04", "CP Decomposition (CP-ALS)", "Execute CP on 3x3x3 tensor with target rank R=2.", "Returns weights vector and 3 factor matrices (3x2).", "PASS"),
        ("TC-05", "Tucker Decomposition (HOOI)", "Execute Tucker on 3x4x5 tensor with ranks [2,2,2].", "Returns 2x2x2 core tensor and 3 factor matrices.", "PASS"),
        ("TC-06", "HOSVD Decomposition", "Execute HOSVD on 4x4x4 tensor.", "Returns orthogonal factor matrices and core tensor.", "PASS"),
        ("TC-07", "Tensor Train Decomposition", "Execute Tensor Train on 3D tensor.", "Returns sequence of 3D TT core tensors.", "PASS"),
        ("TC-08", "Matrix SVD Execution", "Execute SVD on 4x4 matrix input.", "Returns U (4x4), S spectrum (4), V^H (4x4).", "PASS"),
        ("TC-09", "Matrix QR Execution", "Execute QR on 4x3 matrix.", "Returns orthogonal Q (4x4) and upper-triangular R (4x3).", "PASS"),
        ("TC-10", "Matrix LU Execution", "Execute LU decomposition on square matrix.", "Returns lower L and upper U matrices.", "PASS"),
        ("TC-11", "Accuracy Metrics Calculation", "Run Analysis on CP decomposed tensor.", "Calculates MAE, RMSE, Relative Error, and reconstructed head.", "PASS"),
        ("TC-12", "FLOPs & Benchmarking", "Run Benchmark on Tucker decomposition.", "Outputs execution time (ms), FLOPs string, compression ratio.", "PASS"),
        ("TC-13", "Algorithmic Comparison", "Trigger 'Compare' action on 3D tensor input.", "Renders comparative summary table for all algorithms.", "PASS"),
        ("TC-14", "3D Isometric Visualization", "Trigger 'Visualize' action on 3D tensor.", "SVG engine renders 3D scaled block equation diagram.", "PASS"),
        ("TC-15", "Dark/Light Theme Toggle", "Click theme toggle button in header navigation.", "Applies dark/light theme CSS tokens instantly & persists.", "PASS")
    ]

    tbl_tc = doc.add_table(rows=len(test_cases)+1, cols=5)
    tbl_tc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_tc)

    tc_headers = ["Test ID", "Feature Tested", "Test Scenario & Input", "Expected System Output", "Status"]
    c_hdr = tbl_tc.rows[0].cells
    for i, title in enumerate(tc_headers):
        c_hdr[i].width = Inches(0.8 if i==0 else (1.4 if i==1 else (2.0 if i==2 else (1.6 if i==3 else 0.7))))
        set_cell_background(c_hdr[i], "0F172A")
        p = c_hdr[i].paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (t_id, t_feat, t_scen, t_exp, t_stat) in enumerate(test_cases):
        r_cells = tbl_tc.rows[idx+1].cells
        if idx % 2 == 1:
            for c in r_cells: set_cell_background(c, "F8FAFC")
            
        r_cells[0].paragraphs[0].add_run(t_id).font.bold = True
        r_cells[1].paragraphs[0].add_run(t_feat).font.bold = True
        r_cells[2].paragraphs[0].add_run(t_scen)
        r_cells[3].paragraphs[0].add_run(t_exp)
        
        p_stat = r_cells[4].paragraphs[0]
        r_stat = p_stat.add_run(t_stat)
        r_stat.font.bold = True
        r_stat.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # SECTION 8: TIMELINE & GANTT CHART
    # ---------------------------------------------------------------------------
    add_heading_1("8. Timeline")

    add_heading_2("8.1 Project Gantt Chart")
    doc.add_paragraph(
        "The project development timeline spans a 16-week schedule across requirements analysis, algorithm implementation, "
        "frontend visualization engine development, benchmarking integration, testing, and final documentation."
    )

    # Insert Gantt Chart Image
    if os.path.exists("report_assets/gantt_chart.png"):
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img = doc.paragraphs[-1]
        p_img.add_run().add_picture("report_assets/gantt_chart.png", width=Inches(6.2))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figure 8.1: Tensor Decomposer Project Development Gantt Chart (16 Weeks)")
        r_cap.font.bold = True
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = SLATE

    # Draw.io / Mermaid code callout
    gantt_mermaid = """gantt
    title Tensor Decomposer Development Schedule (16 Weeks)
    dateFormat  YYYY-MM-DD
    section Phase 1: Planning
    Requirements Gathering & Proposal       :done, p1, 2026-05-01, 2w
    System Architecture & Math Design       :done, p2, 2026-05-15, 3w
    section Phase 2: Core Engineering
    Core Algorithms (CP, Tucker, HOSVD, TT) :done, p3, 2026-06-05, 4w
    Matrix Operations (SVD, QR, LU, Eig)    :done, p4, 2026-06-19, 2w
    section Phase 3: Frontend & Visuals
    Frontend Development & UI Design        :done, p5, 2026-06-26, 4w
    Interactive SVG & 3D Visualizations     :done, p6, 2026-07-10, 3w
    section Phase 4: Testing & Finalization
    Testing, Benchmarking & Refactoring     :active, p7, 2026-07-24, 3w
    Final Technical Report & Presentation    :active, p8, 2026-08-07, 3w"""
    add_callout_code(gantt_mermaid, "Editable Draw.io / Mermaid Source Code — Gantt Chart Schedule")

    add_heading_2("8.2 Work Breakdown Schedule")
    wbs_data = [
        ("Phase 1: Project Setup & Math Design", "Weeks 1 – 5", "Finalized problem requirements, defined multilinear tensor algorithms (CP-ALS, HOOI), and completed architecture design."),
        ("Phase 2: Backend Algorithm Implementation", "Weeks 4 – 9", "Implemented tensor decomposition algorithms (cp.py, tucker.py, hosvd.py, tensor_train.py) and matrix factorizations."),
        ("Phase 3: Frontend Engine & ES Module Refactor", "Weeks 8 – 13", "Architected responsive UI, built dynamic SVG visualizer (3D isometric diagrams, heatmaps), and refactored home.js into ES modules."),
        ("Phase 4: Analytics, Testing & Deployment", "Weeks 12 – 16", "Integrated FLOPs benchmarking engine, precision error metrics (MAE/RMSE), executed test suite, and authored midterm technical report.")
    ]

    tbl_wbs = doc.add_table(rows=len(wbs_data)+1, cols=3)
    tbl_wbs.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_wbs)

    c_hdr = tbl_wbs.rows[0].cells
    for i, title in enumerate(["Project Phase", "Timeline Schedule", "Key Deliverables & Completed Milestones"]):
        c_hdr[i].width = Inches(2.2 if i==0 else (1.3 if i==1 else 3.0))
        set_cell_background(c_hdr[i], "0F172A")
        p = c_hdr[i].paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (pname, ptime, pdesc) in enumerate(wbs_data):
        r_cells = tbl_wbs.rows[idx+1].cells
        if idx % 2 == 1:
            for c in r_cells: set_cell_background(c, "F8FAFC")
            
        r_cells[0].paragraphs[0].add_run(pname).font.bold = True
        r_cells[1].paragraphs[0].add_run(ptime).font.bold = True
        r_cells[1].paragraphs[0].runs[0].font.color.rgb = TEAL
        r_cells[2].paragraphs[0].add_run(pdesc)

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # SECTION 9 & 10: CONCLUSION & REFERENCES
    # ---------------------------------------------------------------------------
    add_heading_1("9. Conclusion & Future Extensions")
    doc.add_paragraph(
        "The Tensor Decomposer platform successfully addresses the critical need for a unified, accessible, and high-performance system "
        "to analyze, compress, benchmark, and visualize higher-dimensional tensor datasets. By implementing standard multilinear algorithms—including "
        "CP-ALS, HOOI/Tucker-ALS, HOSVD, and Tensor Train Decomposition—alongside traditional matrix factorizations, the platform provides research-grade "
        "decompression capabilities accessible via a modern web interface.\n\n"
        "The integration of single-run computational benchmarking (FLOPs estimation, complexity formulas, parameter compression ratios) and precision reconstruction "
        "error analysis (MAE, RMSE, Relative Error) empowers researchers to make informed trade-offs between compression efficiency and numerical accuracy. "
        "Furthermore, the custom SVG rendering engine delivers intuitive 3D isometric structural equations and component spectra.\n\n"
        "Future Roadmap & Planned Enhancements:\n"
        "1. GPU Acceleration & CUDA Integration: Incorporating PyTorch or CuPy backends to accelerate CP-ALS and HOOI tensor factorizations on large-scale GPU hardware.\n"
        "2. Automated Rank Estimation (CORCONDIA): Implementing automated diagnostic routines (such as CORE Consistency Diagnostic) to guide users in selecting optimal tensor ranks.\n"
        "3. Streaming & Incremental Decompositions: Extending algorithm support to process continuous real-time multiway data streams."
    )

    add_heading_1("10. References")
    references = [
        "1. Kolda, T. G., & Bader, B. W. (2009). Tensor decompositions and applications. SIAM Review, 51(3), 455-500.",
        "2. De Lathauwer, L., De Moor, B., & Vandewalle, J. (2000). A multilinear singular value decomposition. SIAM Journal on Matrix Analysis and Applications, 21(4), 1253-1278.",
        "3. Oseledets, I. V. (2011). Tensor-train decomposition. SIAM Journal on Scientific Computing, 33(5), 2295-2317.",
        "4. Harshman, R. A. (1970). Foundations of the PARAFAC procedure: Models and conditions for an 'explanatory' multimodal factor analysis. UCLA Working Papers in Phonetics, 16, 1-84.",
        "5. Tucker, L. R. (1966). Some mathematical notes on three-mode factor analysis. Psychometrika, 31(3), 279-311.",
        "6. Kiers, H. A. (2000). Towards a standardized notation and terminology in multiway analysis. Journal of Chemometrics, 14(3), 105-122.",
        "7. Cichocki, A., Zdunek, R., Phan, A. H., & Amari, S. I. (2009). Nonnegative matrix and tensor factorizations: applications to exploratory multi-way data analysis and blind source separation. John Wiley & Sons."
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        r.font.size = Pt(9.5)
        r.font.color.rgb = SLATE

    # Add page numbers to headers/footers
    add_header_footer(doc)

    output_path = "SE801_Midterm_Technical_Report_Tensor_Decomposer.docx"
    doc.save(output_path)
    print(f"Report successfully built and saved to {output_path}")

if __name__ == "__main__":
    build_report()
